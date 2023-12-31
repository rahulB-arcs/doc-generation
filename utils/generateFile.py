import os
import glob
from utils.readFile import read_doc, read_csv
from utils.mergeFile import merge_csv
from utils.deleteFile import delete_csv
from docx import Document
from docx.enum.text import WD_BREAK

def generate_file(folder_name, columns_to_save, renamed_columns):

    """
    generate .docx file  with CSV data
    """

    output = []
    doc = Document()
    save_path = './static/generated_docs'

    if not os.path.exists(save_path):
        os.makedirs(save_path)

    merge_csv(folder_name, columns_to_save, renamed_columns)

    csv_path = f'./static/merged_csv/{folder_name}.csv'
    doc_path_list = glob.glob(f'./static/uploads/{folder_name}/docs/*.docx')

    doc_path = doc_path_list[0] if doc_path_list else ''

    text = read_doc(doc_path)
    csv_data = read_csv(csv_path)
    headers = csv_data[0]

    text = replace_text(text)

    rows_to_merge = min(10, len(csv_data) - 1)

    for row in range(1, rows_to_merge + 1):
        values = csv_data[row]
        

        placeholders = { header: value for header, value in zip(headers, values) }
        filled_template = text.format(**placeholders)
        # filled_template = text.format_map({k: f'{{{v}}}' for k, v in placeholders.items()})

        para = doc.add_paragraph(filled_template)
        output.append(para)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    file_path = os.path.join(save_path, f'{folder_name}.docx')
    doc.save(file_path)

    delete_csv(folder_name)

    download_path = f'static/generated_docs/{folder_name}.docx'

    return download_path


def replace_text(text): 
    text = text.replace('{{', '{').replace('}}', '}')
    return text